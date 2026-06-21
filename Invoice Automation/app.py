import os, sys, json, base64, uuid, tempfile, threading
from pathlib import Path
from dotenv import load_dotenv
from flask import Flask, request, jsonify, render_template
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document
from PIL import Image
import io

load_dotenv(Path(__file__).resolve().parent.parent / '.env')

API_KEY = os.getenv("gemini_api_key")
if not API_KEY:
    sys.exit("ERROR: gemini_api_key not found in .env")

client = genai.Client(api_key=API_KEY)
MODEL = "gemini-2.5-flash"

ALLOWED_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".docx"}

DEPARTMENT_MAP = {
    "IT":         ["software", "cloud", "hardware", "license", "subscription", "adobe", "microsoft", "aws", "dell", "saas"],
    "FACILITIES": ["gas", "electricity", "internet", "telephone", "utilities", "energy", "strom", "telekom", "stadtwerke", "e.on", "eon"],
    "FINANCE":    ["consulting", "professional services", "legal", "audit", "advisory", "brightpath"],
    "HR":         ["hotel", "travel", "accommodation", "adlon"],
    "ADMIN":      ["office supplies", "bürobedarf", "stationery", "buerobedarf"],
}

PROMPT = """You are an invoice data extractor. Analyze this invoice and return a JSON object with exactly these fields:
{
  "vendor": "company name",
  "invoice_number": "invoice number or null",
  "date": "invoice date or null",
  "amount_total": "total amount with currency symbol",
  "currency": "EUR or USD etc",
  "invoice_type": "concise type e.g. Gas bill, Software licenses, Cloud services",
  "language": "German or English",
  "summary": "one sentence describing what this invoice is for"
}
Return only valid JSON, no markdown fences."""


def route_department(invoice_type: str, vendor: str) -> str:
    text = ((invoice_type or "") + " " + (vendor or "")).lower()
    for dept, keywords in DEPARTMENT_MAP.items():
        if any(kw in text for kw in keywords):
            return dept
    return "FINANCE"


def analyze_invoice(path: Path) -> dict:
    ext = path.suffix.lower()
    if ext == ".pdf":
        reader = PdfReader(str(path))
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        contents = [PROMPT, f"Invoice text:\n{text}"]
    elif ext in (".png", ".jpg", ".jpeg"):
        img = Image.open(str(path))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        contents = [PROMPT, types.Part.from_bytes(data=buf.getvalue(), mime_type="image/png")]
    elif ext == ".docx":
        doc = Document(str(path))
        lines = []
        for p in doc.paragraphs:
            try:
                t = p.text
                if t and t.strip():
                    lines.append(t)
            except TypeError:
                pass
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    try:
                        t = cell.text
                        if t and t.strip():
                            lines.append(t)
                    except TypeError:
                        pass
        text = "\n".join(lines)
        contents = [PROMPT, f"Invoice text:\n{text}"]
    else:
        raise ValueError(f"Unsupported format: {ext}")

    response = client.models.generate_content(model=MODEL, contents=contents)
    raw = response.text or ""
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


# ── Job state ─────────────────────────────────────────────────────────────────
jobs: dict[str, dict] = {}


def process_paths(job_id: str, paths: list[Path]):
    jobs[job_id]["total"] = len(paths)
    for path in paths:
        try:
            data = analyze_invoice(path)
            def s(key, fallback=""):
                return data.get(key) or fallback
            dept = route_department(s("invoice_type"), s("vendor"))
            jobs[job_id]["results"].append({
                "file": path.name,
                "vendor": s("vendor"),
                "invoice_type": s("invoice_type"),
                "amount": s("amount_total"),
                "date": s("date"),
                "language": s("language"),
                "summary": s("summary"),
                "department": dept,
                "status": "pending",
            })
        except Exception as e:
            jobs[job_id]["results"].append({
                "file": path.name, "status": "error", "error": str(e),
            })
    jobs[job_id]["state"] = "done"


# ── Flask app ──────────────────────────────────────────────────────────────────
app = Flask(__name__)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/process", methods=["POST"])
def process_upload():
    files = request.files.getlist("files")
    if not files or files[0].filename == "":
        return jsonify({"error": "No files uploaded"}), 400

    job_id = uuid.uuid4().hex[:8]
    jobs[job_id] = {"results": [], "state": "processing", "total": len(files)}

    tmp_dir = Path(tempfile.mkdtemp(prefix="invoices_"))
    paths = []
    for f in files:
        p = tmp_dir / f.filename
        f.save(str(p))
        paths.append(p)

    threading.Thread(target=process_paths, args=(job_id, paths), daemon=True).start()
    return jsonify({"job_id": job_id, "total": len(files)})


@app.route("/api/fetch-email", methods=["POST"])
def fetch_email():
    minutes = int(request.json.get("minutes", 10))
    email_addr = os.getenv("gmail_address")
    app_pwd = os.getenv("gmail_app_password")
    if not email_addr or not app_pwd:
        return jsonify({"error": "gmail_address / gmail_app_password missing in .env"}), 400

    job_id = uuid.uuid4().hex[:8]
    jobs[job_id] = {"results": [], "state": "processing", "total": 0}

    def run():
        from email_fetcher import fetch_invoice_attachments
        tmp_dir = Path(tempfile.mkdtemp(prefix="invoices_"))
        paths = fetch_invoice_attachments(email_addr, app_pwd, tmp_dir, last_minutes=minutes)
        process_paths(job_id, paths)

    threading.Thread(target=run, daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/api/status/<job_id>")
def status(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify(job)


@app.route("/api/decision", methods=["POST"])
def decision():
    data = request.json
    job_id = data.get("job_id")
    filename = data.get("file")
    action = data.get("action")  # "approved" or "rejected"
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    for r in job["results"]:
        if r["file"] == filename:
            r["status"] = action
            break
    return jsonify({"ok": True})


if __name__ == "__main__":
    app.run(host="0.0.0.0", debug=True, port=5000)
