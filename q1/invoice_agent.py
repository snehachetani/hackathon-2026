"""
Invoice Processing Agent — Globus Group (St. Wendel)
Reads invoices from a folder, extracts data via Gemini 2.5 Flash,
routes to the right department, and collects one-click approvals.
"""

import os
import sys
import json
import base64
import tempfile
from pathlib import Path
from dotenv import load_dotenv
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document
from PIL import Image
import io

load_dotenv(Path(__file__).resolve().parent.parent / '.env')

API_KEY = os.getenv("gemini_api_key")
if not API_KEY:
    sys.exit("ERROR: gemini_api_key not found in .env")

client = genai.Client(api_key=API_KEY)
MODEL = "gemini-2.5-flash"

DEFAULT_INVOICES_DIR = (
    Path(__file__).resolve().parent.parent / "hackathon_problems_20260620" / "hackathon_problems_20260620" / "questions" / "invoices_hackathon_20260620_part_1"
)

# Parse args: --email to fetch from Gmail, --minutes=N for time window, or a folder path
USE_EMAIL = "--email" in sys.argv
LAST_MINUTES = next(
    (int(a.split("=")[1]) for a in sys.argv[1:] if a.startswith("--minutes=")),
    10,
)
args_without_flags = [a for a in sys.argv[1:] if not a.startswith("--")]
INVOICES_DIR = Path(args_without_flags[0]) if args_without_flags else DEFAULT_INVOICES_DIR

DEPARTMENT_MAP = {
    "IT": ["software", "cloud", "hardware", "license", "subscription", "adobe", "microsoft", "aws", "dell", "saas"],
    "FACILITIES": ["gas", "electricity", "internet", "telephone", "utilities", "energy", "strom", "telekom", "stadtwerke", "e.on", "eon"],
    "FINANCE": ["consulting", "professional services", "legal", "audit", "advisory", "brightpath"],
    "HR": ["hotel", "travel", "accommodation", "adlon"],
    "ADMIN": ["office supplies", "bürobedarf", "stationery", "buerobedarf"],
}


def route_department(invoice_type: str, vendor: str) -> str:
    text = (invoice_type + " " + vendor).lower()
    for dept, keywords in DEPARTMENT_MAP.items():
        if any(kw in text for kw in keywords):
            return dept
    return "FINANCE"


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    lines = []
    for p in doc.paragraphs:
        try:
            t = p.text
            if t and t.strip():
                lines.append(t)
        except TypeError:
            pass
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                try:
                    t = cell.text
                    if t and t.strip():
                        lines.append(t)
                except TypeError:
                    pass
    return "\n".join(lines)


def image_to_base64(path: Path) -> tuple[str, str]:
    img = Image.open(str(path))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode(), "image/png"


PROMPT = """You are an invoice data extractor. Analyze this invoice document and return a JSON object with exactly these fields:
{
  "vendor": "company name",
  "invoice_number": "invoice number or null",
  "date": "invoice date as string or null",
  "amount_total": "total amount with currency symbol",
  "currency": "EUR or USD etc",
  "invoice_type": "concise type e.g. Gas bill, Software licenses, Cloud services, Hardware purchase",
  "language": "German or English",
  "summary": "one sentence describing what this invoice is for"
}
Return only valid JSON, no markdown fences."""


def analyze_invoice(path: Path) -> dict:
    ext = path.suffix.lower()
    contents = []

    if ext == ".pdf":
        text = extract_pdf_text(path)
        contents = [PROMPT, f"Invoice text:\n{text}"]
    elif ext in (".png", ".jpg", ".jpeg"):
        b64, mime = image_to_base64(path)
        contents = [
            PROMPT,
            types.Part.from_bytes(data=base64.b64decode(b64), mime_type=mime),
        ]
    elif ext == ".docx":
        text = extract_docx_text(path)
        contents = [PROMPT, f"Invoice text:\n{text}"]
    else:
        raise ValueError(f"Unsupported format: {ext}")

    response = client.models.generate_content(model=MODEL, contents=contents)
    raw = response.text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


def print_card(filename: str, data: dict, department: str, idx: int, total: int):
    w = 60
    print(f"\n{'='*w}")
    print(f"  Invoice {idx}/{total}: {filename}")
    print(f"{'='*w}")
    print(f"  Vendor:    {data.get('vendor', 'N/A')}")
    print(f"  Type:      {data.get('invoice_type', 'N/A')}")
    print(f"  Amount:    {data.get('amount_total', 'N/A')}")
    print(f"  Date:      {data.get('date', 'N/A')}")
    print(f"  Language:  {data.get('language', 'N/A')}")
    print(f"  Summary:   {data.get('summary', 'N/A')}")
    print(f"  {'-'*(w-2)}")
    print(f"  --> ROUTE TO: {department} DEPARTMENT")
    print(f"{'='*w}")


def fetch_from_email() -> Path:
    from email_fetcher import fetch_invoice_attachments
    email_addr = os.getenv("gmail_address")
    app_pwd = os.getenv("gmail_app_password")
    if not email_addr or not app_pwd:
        sys.exit(
            "ERROR: Set gmail_address and gmail_app_password in .env to use --email mode.\n"
            "  Gmail App Password: https://myaccount.google.com/apppasswords"
        )
    tmp_dir = Path(tempfile.mkdtemp(prefix="invoices_"))
    fetch_invoice_attachments(email_addr, app_pwd, tmp_dir, last_minutes=LAST_MINUTES)
    return tmp_dir


def main():
    scan_dir = INVOICES_DIR
    if USE_EMAIL:
        print("Fetching invoices from Gmail...")
        scan_dir = fetch_from_email()

    invoice_files = sorted([
        f for f in scan_dir.iterdir()
        if f.suffix.lower() in (".pdf", ".png", ".jpg", ".jpeg", ".docx")
    ])

    if not invoice_files:
        sys.exit("No invoice files found in " + str(scan_dir))

    print(f"\n{'#'*60}")
    print(f"  GLOBUS GROUP -- Invoice Processing Agent")
    print(f"  Model: {MODEL}")
    print(f"  Source: {'Gmail inbox' if USE_EMAIL else str(scan_dir)}")
    print(f"  Found {len(invoice_files)} invoices to process")
    print(f"{'#'*60}")

    results = []

    for idx, inv_path in enumerate(invoice_files, 1):
        print(f"\n[{idx}/{len(invoice_files)}] Processing {inv_path.name} ...", end=" ", flush=True)
        try:
            data = analyze_invoice(inv_path)
            dept = route_department(data.get("invoice_type", ""), data.get("vendor", ""))
            print("done")
            print_card(inv_path.name, data, dept, idx, len(invoice_files))

            while True:
                choice = input("  Action: [A]pprove  [R]eject  [S]kip  > ").strip().upper()
                if choice in ("A", "R", "S"):
                    break
                print("  Please enter A, R, or S.")

            status = {"A": "APPROVED", "R": "REJECTED", "S": "SKIPPED"}[choice]
            results.append({
                "file": inv_path.name,
                "vendor": data.get("vendor"),
                "amount": data.get("amount_total"),
                "type": data.get("invoice_type"),
                "department": dept,
                "status": status,
            })
            print(f"  Marked as {status} --> routed to {dept}")

        except Exception as e:
            print(f"ERROR: {e}")
            results.append({"file": inv_path.name, "status": "ERROR", "error": str(e)})

    print(f"\n{'#'*60}")
    print(f"  PROCESSING COMPLETE")
    print(f"{'#'*60}")
    approved = [r for r in results if r["status"] == "APPROVED"]
    rejected = [r for r in results if r["status"] == "REJECTED"]
    print(f"  Approved: {len(approved)}  Rejected: {len(rejected)}  Errors: {len([r for r in results if r['status']=='ERROR'])}")
    print()

    for r in results:
        icon = {"APPROVED": "OK", "REJECTED": "NO", "SKIPPED": "--", "ERROR": "!!"}.get(r["status"], "??")
        dept = r.get("department", "?")
        print(f"  [{icon}] [{dept:<10}] {r['file']:<38} {r.get('amount',''):>12}  {r['status']}")

    out = Path("invoice_results.json")
    with open(out, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    print(f"\n  Results saved to {out}")


if __name__ == "__main__":
    main()
